import logging
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from app.config import Settings
from app.services.cv_storage import CVStorageService
from app.services.pdf_service import SECTION_TITLES, latex_to_plain_text
from app.services.tenant_storage_paths import resolve_tenant_storage_root

logger = logging.getLogger(__name__)


class CVDocxService:
    def __init__(self, settings: Settings, tenant_id: str | None = None) -> None:
        self._settings = settings
        self._tenant_id = tenant_id
        self._tenant_root = resolve_tenant_storage_root(settings, tenant_id)
        self._storage = CVStorageService(settings, tenant_id=tenant_id)

    def generate_tailored_docx(self, project_id: str, job_id: str) -> Path:
        output_dir = self._storage.get_output_dir(project_id, job_id).resolve()
        if not output_dir.exists():
            raise FileNotFoundError(f"Tailored output not found for job {job_id}")
        project = self._storage.get_project(project_id)
        if not project:
            raise FileNotFoundError(f"CV project {project_id} not found")
        title = project.get("name", "Tailored CV")
        return self._build_docx_from_directory(output_dir, title, f"tailored_{job_id[:8]}")

    def generate_master_docx(self, project_id: str) -> Path:
        project = self._storage.get_project(project_id)
        if not project:
            raise FileNotFoundError(f"CV project {project_id} not found")
        master_root = self._storage.get_master_root(project_id).resolve()
        title = project.get("name", "CV")
        return self._build_docx_from_directory(master_root, title, f"master_{project_id[:8]}")

    def _build_docx_from_directory(self, source_dir: Path, title: str, cache_key: str) -> Path:
        document = Document()
        self._configure_document_styles(document)

        contact = self._read_contact_from_resume(source_dir / "resume.tex")
        self._add_header(document, contact or {"name": title, "role": ""})

        sections_dir = source_dir / "sections"
        section_files = sorted(sections_dir.glob("*.tex")) if sections_dir.exists() else []
        if not section_files:
            raise FileNotFoundError("No section files found for DOCX export")

        for section_file in section_files:
            relative = f"sections/{section_file.name}"
            heading = SECTION_TITLES.get(
                relative,
                section_file.stem.replace("_", " ").title(),
            )
            raw_content = section_file.read_text(encoding="utf-8")
            plain_content = latex_to_plain_text(raw_content)
            if not plain_content:
                continue
            self._add_section(document, heading, plain_content)

        docx_dir = self._tenant_root / "reports"
        docx_dir.mkdir(parents=True, exist_ok=True)
        destination = docx_dir / f"cv_{cache_key[:16]}.docx"
        document.save(str(destination))
        logger.info("Generated DOCX at %s", destination)
        return destination

    def _configure_document_styles(self, document: Document) -> None:
        normal_style = document.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(11)

    def _add_header(self, document: Document, contact: dict[str, str]) -> None:
        name = contact.get("name", "CV")
        role = contact.get("role", "")

        name_paragraph = document.add_paragraph()
        name_run = name_paragraph.add_run(name)
        name_run.bold = True
        name_run.font.size = Pt(18)
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if role:
            role_paragraph = document.add_paragraph()
            role_run = role_paragraph.add_run(role)
            role_run.font.color.rgb = RGBColor(61, 90, 128)
            role_run.font.size = Pt(12)
            role_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        contact_bits = [
            value
            for key in ("phone", "city", "email")
            for value in [contact.get(key, "").strip()]
            if value
        ]
        if contact_bits:
            contact_paragraph = document.add_paragraph(" · ".join(contact_bits))
            contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_paragraph()

    def _add_section(self, document: Document, heading: str, content: str) -> None:
        heading_paragraph = document.add_paragraph()
        heading_run = heading_paragraph.add_run(heading.upper())
        heading_run.bold = True
        heading_run.font.size = Pt(12)
        heading_run.font.color.rgb = RGBColor(61, 90, 128)

        bullet_items = self._split_bullet_items(content)
        if bullet_items:
            for item in bullet_items:
                document.add_paragraph(item, style="List Bullet")
        else:
            document.add_paragraph(content)

        document.add_paragraph()

    def _split_bullet_items(self, content: str) -> list[str]:
        if "- " not in content:
            return []
        parts = re.split(r"\s*-\s+", content)
        return [part.strip() for part in parts if part.strip()]

    def _read_contact_from_resume(self, resume_path: Path) -> dict[str, str]:
        if not resume_path.exists():
            return {}

        content = resume_path.read_text(encoding="utf-8")
        contact: dict[str, str] = {}
        field_map = {
            "name": "name",
            "phone": "phone",
            "city": "city",
            "email": "email",
            "LinkedIn": "linkedin",
            "github": "github",
            "role": "role",
        }
        for latex_key, contact_key in field_map.items():
            marker = f"\\def\\{latex_key}{{"
            start_index = content.find(marker)
            if start_index == -1:
                continue
            start_index += len(marker)
            end_index = content.find("}", start_index)
            if end_index == -1:
                continue
            contact[contact_key] = content[start_index:end_index]
        return contact
